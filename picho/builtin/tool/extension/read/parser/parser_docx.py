"""
DOCX parser - converts Word documents to chunks.
"""

import os
import re
from pathlib import Path

from picho.builtin.tool.extension.read.parser.types import (
    Chunk,
    ChunkType,
    Image,
    Metadata,
)


def parse_docx(file_path: str, image_dir: str) -> list[Chunk]:
    """
    Parse a DOCX file and return a list of chunks.

    Args:
        file_path: Path to the DOCX file
        image_dir: Directory to save extracted images

    Returns:
        List of Chunk objects
    """
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    chunks = []
    doc = Document(file_path)

    Path(image_dir).mkdir(parents=True, exist_ok=True)

    image_rels = []
    for i, rel in enumerate(doc.part.rels.values()):
        if rel.reltype == RT.IMAGE:
            image_data = rel.target_part.blob
            image_ext = os.path.splitext(rel.target_ref)[1].lstrip(".")
            image_filename = f"image_{i}.{image_ext}"
            image_path = os.path.join(image_dir, image_filename)

            with open(image_path, "wb") as f:
                f.write(image_data)

            image_rels.append((rel, image_filename))

    body_elements = doc.element.body
    elements_with_positions = []

    for i, paragraph in enumerate(doc.paragraphs):
        elements_with_positions.append((paragraph._p, "paragraph", i))

    for i, table in enumerate(doc.tables):
        elements_with_positions.append((table._tbl, "table", i))

    element_index_map = {elem: idx for idx, elem in enumerate(body_elements)}
    elements_with_positions.sort(
        key=lambda x: element_index_map.get(x[0], float("inf"))
    )

    for element, element_type, index in elements_with_positions:
        if element_type == "paragraph":
            paragraph = doc.paragraphs[index]

            if paragraph.text.strip():
                title_level = _determine_title_level(paragraph)
                chunk = Chunk(
                    type=ChunkType.TEXT.value,
                    text=paragraph.text,
                    image=Image(),
                    metadata=Metadata(
                        is_title=title_level != -1, title_level=title_level
                    ),
                )
                chunks.append(chunk)

            for run in paragraph.runs:
                for rel, image_filename in image_rels:
                    if rel.rId in run._r.xml:
                        chunk = Chunk(
                            type=ChunkType.IMAGE.value,
                            text="",
                            image=Image(type="image_path", path=image_filename),
                        )
                        chunks.append(chunk)
                        image_rels.remove((rel, image_filename))
                        break

        elif element_type == "table":
            table = doc.tables[index]
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append("\t|".join(row_text))

            chunk = Chunk(
                type=ChunkType.TABLE.value,
                text="\n".join(table_text),
                image=Image(),
            )
            chunks.append(chunk)

    return chunks


def _determine_title_level(paragraph) -> int:
    text = paragraph.text.strip()
    if not text:
        return -1

    font_size_pt = _get_paragraph_font_size(paragraph)
    font_size_w = font_size_pt * 2

    is_numbered = _is_numbered_paragraph(paragraph)
    length = len(text)

    if font_size_w >= 48:
        return 1
    elif 36 <= font_size_w < 48:
        return 2
    elif 28 <= font_size_w < 36:
        if length <= 15:
            return 3
        elif is_numbered and length <= 45:
            return 3

    if is_numbered:
        if 28 <= font_size_w < 36 and length <= 45:
            return 3
        elif 24 <= font_size_w < 28 and length <= 45:
            return 4

    return -1


def _get_paragraph_font_size(paragraph) -> float:
    max_size = 0
    for run in paragraph.runs:
        if run.font.size:
            run_size = run.font.size.pt
            if run_size > max_size:
                max_size = run_size
    return max_size


def _is_numbered_paragraph(paragraph) -> bool:
    if paragraph._p.xml.find("w:numPr") != -1:
        return True

    text = paragraph.text.strip()
    if not text:
        return False

    chinese_numbers = [
        "一、",
        "二、",
        "三、",
        "四、",
        "五、",
        "六、",
        "七、",
        "八、",
        "九、",
        "十、",
    ]
    for num in chinese_numbers:
        if text.startswith(num):
            return True

    if text.startswith("第"):
        return True

    if re.match(r"^\d+[.、)]", text):
        return True

    return False
